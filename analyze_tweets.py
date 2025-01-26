import re
import json
import html
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from transformers import pipeline

def create_hyperlink(paragraph, url_text, link_url):
    part = paragraph.part
    r_id = part.relate_to(
        link_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000EE")
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_elm = OxmlElement('w:t')
    text_elm.text = url_text
    new_run.append(text_elm)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def remove_tco_links(text):
    """Remove unnecessary `t.co` links from the tweet."""
    return re.sub(r'https?://t\.co/\S+', '', text)

def fix_double_ellipses(text):
    """Replace redundant ellipses with a single instance."""
    while "... ..." in text:
        text = text.replace("... ...", "...")
    return text

def update_relevance_counters(accounts, username, category):
    """Update the relevance counter for a specific category."""
    if username not in accounts:
        return
    account = accounts[username]
    if "relevance_count" not in account:
        account["relevance_count"] = {"marketing": 0, "AI": 0, "Crypto": 0}
    account["relevance_count"][category] += 1
    account["last_checked"] = datetime.utcnow().isoformat()

def save_accounts_relevant(accounts):
    """Save the updated accounts relevance data."""
    with open("data/accounts_relevant.json", "w", encoding="utf-8") as f:
        json.dump(list(accounts.values()), f, indent=2, ensure_ascii=False)

def main():
    try:
        with open("all_tweets.json", "r", encoding="utf-8") as f:
            all_data = json.load(f)
    except FileNotFoundError:
        print("ERROR: all_tweets.json not found. Run fetch_tweets.py first.")
        return

    try:
        with open("data/accounts_relevant.json", "r", encoding="utf-8") as f:
            accounts_relevant = {acc["username"]: acc for acc in json.load(f)}
    except FileNotFoundError:
        print("ERROR: accounts_relevant.json not found. Ensure it exists.")
        return

    classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
    THRESHOLD = 0.90
    LABELS = ["marketing", "AI", "Crypto"]

    marketing_posts = []
    ai_posts = []
    crypto_posts = []
    total_count = 0

    for username, tweets in all_data.items():
        for tw in tweets:
            total_count += 1

            text = tw.get("text", "").strip()
            if not text:
                continue

            # Perform relevance check first
            result = classifier(
                sequences=text,
                candidate_labels=LABELS,
                multi_label=True,
                hypothesis_template="This tweet is about {}."
            )
            label_scores = dict(zip(result["labels"], result["scores"]))

            # Skip irrelevant tweets
            if all(score < THRESHOLD for score in label_scores.values()):
                continue

            # If relevant, proceed with cleaning and formatting
            raw_date = tw.get("created_at", "")
            t_id = tw.get("id", "")
            if not raw_date or not t_id:
                continue

            try:
                dt_obj = datetime.strptime(raw_date, "%Y-%m-%dT%H:%M:%S.%fZ")
                date_str = dt_obj.strftime("%A, %B %d, %Y @ %I:%M:%S %p UTC")
            except ValueError:
                date_str = raw_date

            text = html.unescape(text)
            text = remove_tco_links(text)
            text = fix_double_ellipses(text)

            final_link = f"https://x.com/{username}/status/{t_id}"  # Update domain

            # Add to the respective category
            if label_scores.get("marketing", 0.0) >= THRESHOLD:
                marketing_posts.append({"username": username, "text": text, "date_str": date_str, "tweet_id": t_id})
                update_relevance_counters(accounts_relevant, username, "marketing")

            if label_scores.get("AI", 0.0) >= THRESHOLD:
                ai_posts.append({"username": username, "text": text, "date_str": date_str, "tweet_id": t_id})
                update_relevance_counters(accounts_relevant, username, "AI")

            if label_scores.get("Crypto", 0.0) >= THRESHOLD:
                crypto_posts.append({"username": username, "text": text, "date_str": date_str, "tweet_id": t_id})
                update_relevance_counters(accounts_relevant, username, "Crypto")

    save_accounts_relevant(accounts_relevant)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for heading, posts in [
        ("Original Marketing Posts from Small Accounts", marketing_posts),
        ("Original AI Posts from Small Accounts", ai_posts),
        ("Original Crypto Posts from Small Accounts", crypto_posts),
    ]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph("")  # Add space after the header
        if not posts:
            doc.add_paragraph(f"No relevant {heading.split()[1]} tweets found.\n")
        else:
            for item in posts:
                head_p = doc.add_paragraph()
                head_p.add_run(f"@{item['username']}").bold = True
                head_p.add_run(f" | {item['date_str']}").bold = True

                body_p = doc.add_paragraph()
                body_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                body_p.add_run(item["text"])

                vt_p = doc.add_paragraph()
                vt_p.add_run("View tweet: ").italic = True
                create_hyperlink(vt_p, final_link, final_link)

                doc.add_paragraph("")  # Add space between tweets

    doc.save("all_tweets_relevant.docx")
    print(f"Done! Processed {total_count} tweets. Results saved in 'all_tweets_relevant.docx'.")

if __name__ == "__main__":
    main()
